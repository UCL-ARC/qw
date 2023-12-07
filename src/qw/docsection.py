"""
DocX outline level iterator.

Layer over python-docx allowing looping through
outline sections, duplicating them as necessary
to put the required data in.
"""
import re
from copy import deepcopy
from typing import Optional, TypeAlias

import docx
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from loguru import logger
from lxml.etree import QName

from qw import md
from qw.base import QwError


def qname(ns: str, val: str) -> str:
    """Qualified name for "ns:val"."""
    return QName(docx.oxml.ns.nsmap[ns], val)


def namespace_w(val: str) -> str:
    """Qualified name for "w:val"."""
    return qname("w", val)


MERGEFIELD_RE = re.compile(r"\s*MERGEFIELD\s+(.*?)\s*")
ATTRIB_ABSTRACTNUMID = namespace_w("abstractNumId")
ATTRIB_NUMID = namespace_w("numId")
ATTRIB_VAL = namespace_w("val")
ATTRIB_ASCII = namespace_w("ascii")
ATTRIB_HANSI = namespace_w("hAnsi")
ATTRIB_SPACE = qname("xml", "space")
TAG_R = namespace_w("r")
TAG_RPR = namespace_w("rPr")
TAG_B = namespace_w("b")
TAG_I = namespace_w("i")
TAG_RFONTS = namespace_w("rFonts")
TAG_T = namespace_w("t")
TAG_BR = namespace_w("br")
TAG_R_ID = qname("r", "id")
TAG_RSTYLE = namespace_w("rStyle")
TAG_P = namespace_w("p")
TAG_PPR = namespace_w("pPr")
TAG_PSTYLE = namespace_w("pStyle")
TAG_NUMPR = namespace_w("numPr")
TAG_ILVL = namespace_w("ilvl")
TAG_NUMID = namespace_w("numId")
TAG_BIDI = namespace_w("bidi")
TAG_JC = namespace_w("jc")
TAG_HYPERLINK = namespace_w("hyperlink")


def is_field_beginning(run):
    """Is this <w:r> run element the beginning of a field?."""
    return len(run.xpath('w:fldChar[@w:fldCharType="begin"]')) > 0


def is_field_separator(run):
    """
    Is this <w:r> run element a field separator?.

    There is a separator between a field's metadata and its text.
    """
    return len(run.xpath('w:fldChar[@w:fldCharType="separate"]')) > 0


def is_field_ending(run):
    """Is this <w:r> run element the end of a field?."""
    return len(run.xpath('w:fldChar[@w:fldCharType="end"]')) > 0


HEADING_STYLE_ID = [
    "Heading1",
    "Heading2",
    "Heading3",
    "Heading4",
]

PREFORMATTED_FONT_NAME = "Courier"

_DocSection: TypeAlias = "DocSection"

# Paragraph level (in DocSection) for a list item
_LIST_ITEM_BASE_LEVEL = 50
# Paramgraph level (in DocSection) for a heading
_HEADING_BASE_LEVEL = -50


class DocSection:
    """
    A two-dimensional iteration through an MS Word document.

    The iteration refers to an outline section; a heading together
    with all the paragraphs beneath it, including those headings of
    deeper levels, or a paragraph or list item together with all the
    more indented list items directly below it.

    One direction of iteration is with the `deeper` method. This
    returns a new DocSection that iterates through the paragraphs in
    this section but at a deeper outline level. If we are only
    looking at one paragraph, `deeper` will return None as no deeper
    iteration is possible.

    The other direction of iteration is with the `next` method. This
    moves the iteration to the next section in the document. If this
    iterator was created with the `deeper` method `next` will not
    iterate beyond the bounds of the section the original iteration
    was pointing to; in this case, or if the iteration hits the end
    of the document, `next` returns None.
    """

    def __init__(self, docx, from_index=None, outer_level=None, parent=None):
        """
        Create a DocSection as an empty object.

        `next` should be called to initialize.
        The section begins after the `after` element if set, or at
        the start of the document if not.
        The iteration ends at the end of the document or just before
        the first paragraph that has a level equal to or lower than
        outer_level (if not None)
        """
        self.docx = docx
        bodies = docx.element.xpath("w:body")
        if len(bodies) != 1:
            msg = "not a single body"
            raise QwError(msg)
        # The "w:body" element of the document
        self.element = bodies[0]
        # The index of the first paragraph pointed to (within the
        # body element)
        self.start_index = from_index or 0
        # The index just past the last paragraph pointed to
        self.end_index = self.start_index
        # The outline level (header level or list item indent)
        # of the section pointed to.
        self.start_level = None
        # The outline level of the section that created this with
        # its `deeper` method.
        self.outer_level = outer_level
        # The DocSection that created this DocSection with its
        # `deeper` method.
        self.parent = parent
        # The paragraph after the current section and all of its
        # duplicates, or None if it is the end of the document.
        self.iteration_end = self.element[self.start_index]

    def first_paragraph_text(self) -> str | None:
        """Get the text in the first paragraph."""
        if self.end_index == self.start_index:
            return None
        return self.element[self.start_index].text

    def _get_number_level(self, element):
        """
        Get the number level.

        This is found either in the paragraph formatting
        or the number style, and refers to the indent level
        of a list item.

        We return the indent level plus _LIST_ITEM_BASE_LEVEL;
        list items are deeper than body text (at level 0).
        """
        bullets = element.xpath("w:pPr/w:numPr/w:ilvl")
        if len(bullets) != 0:
            val = bullets[0].attrib[ATTRIB_VAL]
            if val.isnumeric():
                return int(val) + _LIST_ITEM_BASE_LEVEL
            return _LIST_ITEM_BASE_LEVEL
        return None

    def _get_paragraph_level(self, paragraph):
        """
        Get the outline level of the paragraph in `element`.

        `element` is the <w:p> element in the document. The outline
        level of body text is 0. Headings are shallower, at
        _HEADING_BASE_LEVEL (for Heading 1) counting up.
        """
        # Find the style outline level element
        styles = paragraph.xpath("w:pPr/w:pStyle")
        if len(styles) != 0:
            style_id = styles[0].attrib[ATTRIB_VAL]
            style = self.docx.styles.get_by_id(
                style_id,
                docx.enum.style.WD_STYLE_TYPE.PARAGRAPH,
            )
            outlines = style.element.xpath("w:pPr/w:outlineLvl")
            if len(outlines) != 0:
                val = outlines[0].attrib[ATTRIB_VAL]
                if val.isnumeric():
                    # We have found the heading level in the style
                    return int(val) + _HEADING_BASE_LEVEL
            # If the style does not have an outline level, maybe
            # it has a number style with its own level.
            r = self._get_number_level(style.element)
            if r is not None:
                return r
        # bullet level in the paragraph format itself.
        r = self._get_number_level(paragraph)
        if r is not None:
            return r
        return 0

    def _depth(self):
        if self.parent is None:
            return 0
        return 1 + self.parent._depth()

    def at_iteration_end(self) -> bool:
        """Determine if we are on the last iteration."""
        if self.at_document_end():
            return True
        return self.iteration_end == self.element[self.end_index]

    def next_section(self) -> bool:
        """
        Move to the next section in the document.

        :return: True if we have advanced to the next section, or
        False if we did not because we are at the end of this level.
        """
        if self.at_document_end():
            return False
        self.start_index = self.end_index
        p = self.element[self.start_index]
        level = self._get_paragraph_level(p)
        if self.outer_level is not None and level <= self.outer_level:
            # We have finished this run of sections at this level
            return False
        if self.start_level is None or level < self.start_level:
            self.start_level = level
        # We start with a head paragraph at start_index. Now we
        # will determine how many subsequent paragraphs also belong
        # to this section.
        p = p.getnext()
        # For now we only have one head paragraph, but later we might
        # have more if its text gets replaced.
        self.last_head_paragraph_index = self.start_index
        self.end_index += 1
        while p is not None:
            level = self._get_paragraph_level(p)
            # Have we reached the end of the section?
            if level <= self.start_level:
                logger.debug(
                    "Section {0} - {1} (hit level {2} <= {3}). Depth {4}",
                    self.start_index,
                    self.end_index,
                    level,
                    self.start_level,
                    self._depth(),
                )
                if self.iteration_end == self.element[self.start_index]:
                    self.iteration_end = self.element[self.end_index]
                return True
            self.end_index += 1
            p = p.getnext()
        logger.debug(
            "Section {0} - {1} (end of document). Depth {2}",
            self.start_index,
            self.end_index,
            self._depth(),
        )
        self.iteration_end = None
        return True

    def at_document_end(self):
        """Have we reached the document end?."""
        return len(self.element) <= self.end_index

    def deeper(self) -> Optional[_DocSection]:  # noqa: UP007
        """
        Get a deeper iterator.

        Returns a new DocSection that iterates through the current
        section at a deeper level, if possible.
        """
        if self.end_index - self.start_index <= 1:
            # We only have one paragraph, so there is no deeper
            # iteration to be done.
            return None
        return DocSection(
            self.docx,
            from_index=self.start_index + 1,
            outer_level=self.start_level,
            parent=self,
        )

    def duplicate(self):
        """
        Add a copy of this section after it.

        Note that any fields that have already been substituted will
        have this same substitution in the duplicate.
        """
        length = self.end_index - self.start_index
        # Copy the paragraphs one-by-one
        for i in range(length):
            p = deepcopy(self.element[self.start_index + i])
            self.element.insert(self.end_index + i, p)
        if self.parent is not None:
            # We need to tell the shallower iteration that
            # it now has more paragraphs in its iteration.
            self.parent._paragraph_count_changed(self.end_index, length)

    def _paragraph_count_changed(self, start: int, length: int):
        """
        Update the paragraph count.

        This is called when a DocSection created by `deeper` either
        changes the number of paragraphs it has or hears of changes
        that happened even deeper.

        start -- the index of the first new paragraph that was added
        length -- the number of new paragraphs
        """
        if start <= self.start_index:
            # we are adding paragraphs before this section
            self.start_index += length
        elif start <= self.end_index:
            logger.debug(
                "Adding {} paragraphs to depth {}",
                length,
                self._depth(),
            )
            self.end_index += length
        if self.parent is not None:
            self.parent._paragraph_count_changed(start, length)

    def fields(self) -> set[str]:
        """
        Find the fields in the first paragraph.

        Returns a set of all the fields present in the first
        paragraph of this section.
        """
        r: set[str] = set()
        if self.end_index == self.start_index:
            return r
        for instr in self.element[self.start_index].xpath("*/w:instrText"):
            merge = MERGEFIELD_RE.fullmatch(instr.text)
            if merge is not None:
                r.add(merge.group(1))
        return r

    def paragraph_is_only_field(self) -> bool:
        """
        Is this paragraph only a field?.

        Returns True if there is exactly one field in the first
        paragraph of this section (with no other text) and it is
        normal body text.

        If this is true, the whole paragraph should be replaced with
        rich text. If not, each field is replaced with plain text.
        """
        if self.end_index == self.start_index:
            return False
        paragraph = self.element[self.start_index]
        if self._get_paragraph_level(paragraph) != 0:
            # Not normal body text, so shouldn't be replaced with
            # rich text.
            return False
        is_in_field = False
        for run in paragraph.xpath("w:r"):
            if is_field_beginning(run):
                # We are in a field (which spans multiple <w:r>s)
                is_in_field = True
            else:
                if not is_in_field:
                    # We have found a <w:r> not in a field.
                    return False
                if is_field_ending(run):
                    is_in_field = False
        # There is no run that is not part of a field.
        return True

    def remove_nonfirst_paragraphs(self):
        """
        Remove all but the first paragraph of this section.

        This is used when we want to replace the whole section with
        a single "None" text.
        """
        start = self.start_index + 1
        length = self.end_index - start
        if length <= 0:
            return
        for _i in range(length):
            self.element.remove(self.element[start])
        self.end_index = start
        if self.parent:
            self.parent._paragraph_count_changed(start, -length)

    def replace_first_paragraph(
        self,
        typ: md.DocumentBuilder.ParagraphType,
        level: int = 0,
    ):
        """
        Replace the first paragraph with an empty paragraph.

        Replaces the first paragraph in the section with an empty
        paragraph of the required type.

        `level` is the level of the heading (0 = Heading 1)
        or bullet/numbered list element (0 = top)
        """
        p = self.element[self.start_index]
        p.clear()
        self._add_paragraph_style(typ, level, p)

    def _get_num_id(self, fmt):
        """
        Get the numId required for the bullet style we want.

        We search the styles for a style that has the required format
        at one of its levels. It seems as though all styles have the
        same format at each level.
        We are assuming that we only need any arbitrary one.
        fmt -- the bullet format we want, decimal, bullet or maybe
            something else.
        """
        numbering = self.docx.part.numbering_part.element
        num_fmts = numbering.xpath(
            f'w:abstractNum/w:lvl/w:numFmt[@w:val="{fmt}"]',
        )
        if len(num_fmts) == 0:
            return None
        abstract_num_id = (
            num_fmts[0].getparent().getparent().attrib[ATTRIB_ABSTRACTNUMID]
        )
        nums = numbering.xpath(
            f'w:num/w:abstractNumId[@w:val="{abstract_num_id}"]',
        )
        if len(nums) == 0:
            return None
        return nums[0].getparent().attrib[ATTRIB_NUMID]

    def _add_paragraph_style(self, typ, level, paragraph):
        """
        Apply a style to the paragraph.

        Applies the DocumentBuilder paragraph format specifiers to
        a paragraph.
        """
        ppr = paragraph.makeelement(TAG_PPR)
        paragraph.append(ppr)
        num_id = None
        style_name = None
        if typ is None:
            # Body text
            style_name = "Normal"
        elif typ == md.DocumentBuilder.ParagraphType.HEADING:
            # Heading
            if level >= 0 and level < len(HEADING_STYLE_ID):
                style_name = HEADING_STYLE_ID[level]
        elif typ == md.DocumentBuilder.ParagraphType.LIST_ORDERED:
            # Numbered list item
            style_name = "Normal"
            num_id = self._get_num_id("decimal")
        elif typ == md.DocumentBuilder.ParagraphType.LIST_UNORDERED:
            # Bullet list item
            style_name = "Normal"
            num_id = self._get_num_id("bullet")
        if style_name is not None:
            ppr.append(
                paragraph.makeelement(TAG_PSTYLE, {ATTRIB_VAL: style_name}),
            )
        if num_id is not None:
            num_pr = paragraph.makeelement(TAG_NUMPR)
            num_pr.extend(
                [
                    paragraph.makeelement(TAG_ILVL, {ATTRIB_VAL: str(level)}),
                    paragraph.makeelement(TAG_NUMID, {ATTRIB_VAL: num_id}),
                ],
            )
            ppr.append(num_pr)
        # Basic stuff that's in all paragraph formats
        ppr.extend(
            [
                # Bidirectional information
                paragraph.makeelement(TAG_BIDI, {ATTRIB_VAL: "0"}),
                # Justification
                paragraph.makeelement(TAG_JC, {ATTRIB_VAL: "left"}),
                # Standard run format
                paragraph.makeelement(TAG_RPR),
            ],
        )

    def add_run(
        self,
        text: str,
        bold: md.DocumentBuilder.Boldness | None = None,
        italic: md.DocumentBuilder.Italicness | None = None,
        pre: md.DocumentBuilder.Spacing | None = None,
    ):
        """
        Add a run to the end of the head of the current section.

        DocBuilder override.
        """
        if not text:
            return
        p = self.element[self.last_head_paragraph_index]
        r = p.makeelement(TAG_R)
        p.append(r)
        rpr = p.makeelement(TAG_RPR)
        if bold:
            rpr.extend(rpr.makeelement(TAG_B))
        if italic:
            rpr.append(rpr.makeelement(TAG_I))
        if pre:
            rpr.append(
                rpr.makeelement(
                    TAG_RFONTS,
                    {
                        ATTRIB_ASCII: PREFORMATTED_FONT_NAME,
                        ATTRIB_HANSI: PREFORMATTED_FONT_NAME,
                    },
                ),
            )
        r.append(rpr)
        self._add_plain_text(r, text)

    def _add_plain_text(self, r, text):
        """
        Add a plain text run.

        Adds plain text to a <w:r> element as a series of text and
        linebreak elements.
        """
        br = False
        for text_line in text.splitlines():
            if br:
                r.append(r.makeelement(TAG_BR))
            else:
                br = True
            t = r.makeelement(TAG_T, {ATTRIB_SPACE: "preserve"})
            t.text = text_line
            r.append(t)

    def _get_link_id(self, link: str):
        """
        Get a link ID for the target of a hyperlink.

        MS Word documents store link targets separately from the
        text. We must search the links part for the target to see
        if it already exists. If it does, we return its ID, if not
        we create a new one and return its ID.
        """
        # find an existing link
        return self.docx.part.relate_to(link, RT.HYPERLINK, is_external=True)

    def add_hyperlink(self, text: str, link: str):
        """
        Add a hyperlink.

        DocBuilder override adds a hyperlink to the end of the
        current section.
        """
        p = self.element[self.last_head_paragraph_index]
        rid = self._get_link_id(link)
        hyperlink = p.makeelement(TAG_HYPERLINK, {TAG_R_ID: rid})
        run = p.makeelement(TAG_R)
        rpr = p.makeelement(TAG_RPR)
        r_style = p.makeelement(TAG_RSTYLE, {ATTRIB_VAL: "InternetLink"})
        rpr.append(r_style)
        # Make the actual visible text part
        t = p.makeelement(TAG_T)
        t.text = text
        run.extend([rpr, t])
        hyperlink.append(run)
        p.append(hyperlink)

    def add_paragraph(self, typ: md.DocumentBuilder.ParagraphType, level: int = 0):
        """
        Add an empty paragraph of the required style.

        Adds a new empty paragraph to the end of the head of the
        current section.

        DocBuilder override.
        """
        p = self.element.makeelement(TAG_P)
        self._add_paragraph_style(typ, level, p)
        self.last_head_paragraph_index += 1
        self.element.insert(self.last_head_paragraph_index, p)

    def _delete_backwards_until(self, node, predicate):
        """
        Delete nodes backwards.

        Deletes XML sibling nodes backwards until finding one that
        matches the predicate.

        Returns the node before the deleted nodes, if there is one.

        Both the starting node and the matching node are deleted.
        A node matches the predicate if predicate(node) returns True.
        """

        def go_next(n):
            return n.getprevious()

        return self._delete_direction_until(node, predicate, go_next)

    def _delete_forwards_until(self, node, predicate):
        """
        Delete node forwards.

        Deletes XML sibling nodes forwards until finding one that
        matches the predicate.

        Returns the node after the deleted nodes, if there is one.

        Both the starting node and the matching node are deleted.
        A node matches the predicate if predicate(node) returns True.
        """

        def go_next(n):
            return n.getnext()

        return self._delete_direction_until(node, predicate, go_next)

    def _delete_direction_until(self, node, predicate, go_next):
        """
        Delete nodes backwards or forwards.

        Delete XML nodes until finding one that matches the
        predicate.

        Returns the next node after the deleted nodes, if there is
        one.

        The direction of travel is given by the go_next function that
        takes and returns an XML node.
        Both the starting node and the matching node are deleted.
        A node matches the predicate if predicate(node) returns True.
        """
        while node is not None:
            next_node = go_next(node)
            at_end = predicate(node)
            node.getparent().remove(node)
            if at_end:
                return next_node
            node = next_node
        return None

    def _replace_field_runs(self, instr_node, plain_text: str):
        """
        Replace the fields containing instr_node with plain_text.

        instr_node -- the <w:instrText> node within the field
        plain_text -- replacement text
        """
        # Find the start of the field
        r = instr_node.getparent()
        self._delete_backwards_until(r.getprevious(), is_field_beginning)
        # Find the normal text node
        r = self._delete_forwards_until(r, is_field_separator)
        if r is not None:
            # Remove everything from this run
            for t in r.xpath("w:t"):
                r.remove(t)
            # Add our new text to this run
            self._add_plain_text(r, plain_text)
            # And delete the rest of the field furniture
            self._delete_forwards_until(r.getnext(), is_field_ending)

    def replace_field(self, name: str, plain_text: str) -> int:
        """
        Replace all MailMerge fields named `name` with pain text.

        Returns the number of fields thus replaced.
        """
        count = 0
        for instr in self.element[self.start_index].xpath("*/w:instrText"):
            merge = MERGEFIELD_RE.fullmatch(instr.text)
            if merge is not None and merge.group(1) == name:
                self._replace_field_runs(instr, plain_text)
                count += 1
        return count


class DocSectionParagraphReplacer(md.DocumentBuilder):
    """
    DocumentBuilder that replaces a single paragraph.

    The paragraph in question is the first paragraph pointed to
    by a DocSection object.
    """

    def __init__(self, section: DocSection):
        """Initialize with a DocSection."""
        self.section = section
        self.first = True

    def new_paragraph(
        self,
        paragraph_type: md.DocumentBuilder.ParagraphType = None,
        paragraph_level: int = 0,
    ):
        """DocumentBuilder override."""
        if self.first:
            self.section.replace_first_paragraph(paragraph_type, paragraph_level)
            self.first = False
        else:
            self.section.add_paragraph(paragraph_type, paragraph_level)

    def add_run(
        self,
        text: str,
        bold: md.DocumentBuilder.Boldness | None = None,
        italic: md.DocumentBuilder.Italicness | None = None,
        pre: md.DocumentBuilder.Spacing | None = None,
    ):
        """DocumentBuilder override."""
        self.section.add_run(text, bold, italic, pre)

    def add_hyperlink(
        self,
        text: str,
        link: str,
    ):
        """DocumentBuilder override."""
        self.section.add_hyperlink(text, link)

    def end(self):
        """DocumentBuilder override."""
