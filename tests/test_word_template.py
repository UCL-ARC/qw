"""Replacing fields with data; high-level test."""
from itertools import chain
from tempfile import NamedTemporaryFile

import docx

from qw.mergedoc import load_template


def concat(list_of_lists):
    """Return flattened list of lists."""
    return list(chain.from_iterable(list_of_lists))


def test_replace_fields_with_data():
    """Test replacing fields through MergeData."""
    doc = load_template(
        "tests/resources/msword/DocSection_fields.docx",
    )
    with NamedTemporaryFile("w+b") as temp_file:
        data = {
            "software-component": [
                {
                    "id": 12,
                    "name": "twelve",
                    "description": "A one then a two is a twelve.",
                },
                {
                    "id": 25,
                    "name": "twenty-five",
                    "description": "A quarter century.",
                },
                {
                    "id": 100,
                    "name": "one-hundred",
                    "description": "The ton!",
                },
            ],
            "soup": [
                {
                    "id": 42,
                    "name": "python",
                    "description": "Dynamically typed language with batteries included.",
                },
                {
                    "id": 43,
                    "name": "qw",
                    "description": "Regulation checking and documentation tool.",
                },
            ],
        }
        doc.write(temp_file, data)
        dx = docx.Document(temp_file.name)
        pwf_para = "Paragraph with fields {id} and {name}."
        expecteds = [
            "Heading One",
            *concat(
                [
                    [pwf_para.format(**sc), sc["description"]]
                    for sc in data["software-component"]
                ],
            ),
            "Heading Two",
            "Paragraph C",
            "Unordered list",
            "With some",
            "indents",
            "and more",
            "bullets",
            "Second heading two",
            *concat(
                [
                    ["Soup item {}".format(soup["name"]), soup["description"]]
                    for soup in data["soup"]
                ],
            ),
            "Another top-level heading",
            "Paragraph E",
            "",
        ]
        for p, expected in zip(dx.paragraphs, expecteds, strict=True):
            assert p.text == expected


def test_replace_hierarchical_fields_with_data():
    """Test replacing hierarchical fields through MergeData."""
    doc = load_template(
        "tests/resources/msword/DocSection_two_level_fields.docx",
    )
    with NamedTemporaryFile("w+b") as temp_file:
        data = {
            "software-requirement": [
                {
                    "id": 1,
                    "name": "Display dose",
                    "description": "Put the dose on the screen.",
                    "system-requirement": 101,
                },
                {
                    "id": 2,
                    "name": "Update dose",
                    "description": "The dose on screen is updated as the controls are adjusted.",
                    "system-requirement": 101,
                },
                {
                    "id": 3,
                    "name": "Stop plunger",
                    "description": "Stop the plunger as the required dose is met",
                    "system-requirement": 102,
                },
            ],
            "system-requirement": [
                {
                    "id": 101,
                    "name": "Show the dose",
                    "description": "The user needs to be able to read out the dose on a screen.",
                },
                {
                    "id": 102,
                    "name": "Deliver the dose",
                    "description": "The patient needs to receive the expected dose.",
                },
                {
                    "id": 103,
                    "name": "Have a party",
                    "description": "The developers need a break.",
                },
            ],
        }
        doc.write(temp_file, data)
        dx = docx.Document(temp_file.name)
        expecteds = []
        for sysreq in data["system-requirement"]:
            sysr_id = sysreq["id"]
            expecteds.extend(
                [
                    f"System requirement {sysr_id}",
                    sysreq["name"],
                    sysreq["description"],
                ],
            )
            softreqs = list(
                filter(
                    lambda soft: soft["system-requirement"] == sysr_id,
                    data["software-requirement"],
                ),
            )
            expecteds.append("Implemented by the following software requirements:")
            if len(softreqs) == 0:
                expecteds.append("None.")
            else:
                for softreq in softreqs:
                    expecteds.extend(
                        [
                            "{id}: {name}".format(**softreq),
                            softreq["description"],
                        ],
                    )

        expecteds.extend(
            [
                "Afterword",
                "Some text here.",
                "",
            ],
        )
        for p, expected in zip(dx.paragraphs, expecteds, strict=True):
            assert p.text == expected
