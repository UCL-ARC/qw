"""Tests for DocSection."""

import docx

from qw.docsection import DocSection


def test_docsection_iteration():
    """Test docstring iteration works."""
    doc = docx.Document(
        "tests/resources/msword/DocSection_no_fields.docx",
    )
    section = DocSection(doc)
    v = section.next_section()
    assert v, "Failed to find the first section"
    assert section.first_paragraph_text() == "Heading One"
    heading_one_section = section.deeper()
    v = section.next_section()
    assert v, "Failed to find the second section"
    assert section.first_paragraph_text() == "Another top-level heading"
    v = section.next_section()
    assert not v, "Did not expect a third section"
    # deeper section
    heading_one_section.next_section()
    assert heading_one_section.first_paragraph_text() == "Paragraph A"
    heading_one_section.next_section()
    assert heading_one_section.first_paragraph_text() == "Paragraph B"
    heading_one_section.next_section()
    assert heading_one_section.first_paragraph_text() == "Heading Two"
    heading_two_section = heading_one_section.deeper()
    heading_one_section.next_section()
    assert heading_one_section.first_paragraph_text() == "Second heading two"
    v = heading_one_section.next_section()
    assert not v, "Did not expect a deeper section after Second heading two"
    # even deeper section
    v = heading_two_section.next_section()
    assert heading_two_section.first_paragraph_text() == "Paragraph C"
    paragraph_c_section = heading_two_section.deeper()
    v = heading_two_section.next_section()
    assert not v, "Did not expect a section after Paragraph C"
    paragraph_c_section.next_section()
    assert paragraph_c_section.first_paragraph_text() == "Unordered list"
    paragraph_c_section.next_section()
    assert paragraph_c_section.first_paragraph_text() == "With some"
    indented_section = paragraph_c_section.deeper()
    indented_section.next_section()
    assert indented_section.first_paragraph_text() == "indents"
    indented_section.next_section()
    assert indented_section.first_paragraph_text() == "and more"
    v = indented_section.next_section()
    assert not v, "Did not expect a section after ordered list"


def section_is_as_expected(
    section: DocSection,
    expecteds: list[str | list],
):
    """Test section has the text we expect."""
    for e in expecteds:
        if isinstance(e, list):
            section_is_as_expected(section.deeper(), e)
        else:
            section.next_section()
            assert section.first_paragraph_text() == e


def test_docsection_duplication():
    """Test duplication of DocSections."""
    doc = docx.Document(
        "tests/resources/msword/DocSection_no_fields.docx",
    )
    s1 = DocSection(doc)
    s1.next_section()
    s2 = s1.deeper()
    s2.next_section()
    s2.next_section()
    s2.next_section()
    assert s2.first_paragraph_text() == "Heading Two"
    s3 = s2.deeper()
    s3.next_section()
    s4 = s3.deeper()
    assert s4 is not None
    s4.next_section()
    s4.next_section()
    assert s4.first_paragraph_text() == "With some"
    s4.duplicate()
    s5 = s4.deeper()
    s5.next_section()
    assert s5.first_paragraph_text() == "indents"
    s5.duplicate()
    s4.duplicate()
    assert s3.first_paragraph_text() == "Paragraph C"
    s4 = s3.deeper()
    assert s4 is not None
    section_is_as_expected(
        s3.deeper(),
        [
            "Unordered list",
            "With some",
            ["indents", "indents", "and more"],
            "With some",
            ["indents", "indents", "and more"],
            "With some",
            ["indents", "and more"],
            "bullets",
        ],
    )


def test_docsection_delete_nonhead():
    """Test deletion of paragraphs in DocSections."""
    doc = docx.Document(
        "tests/resources/msword/DocSection_no_fields.docx",
    )
    s1 = DocSection(doc)
    s1.next_section()
    s2 = s1.deeper()
    s2.next_section()
    s2.next_section()
    s2.next_section()
    s3 = s2.deeper()
    s3.next_section()
    s3.remove_nonfirst_paragraphs()
    assert s3.first_paragraph_text() == "Paragraph C"
    assert s3.deeper() is None
    assert not s3.next_section()
    s2.next_section()
    assert s2.first_paragraph_text() == "Second heading two"


def test_docsection_replace_paragraph():
    """Test paragraph replacement."""
    doc = docx.Document(
        "tests/resources/msword/DocSection_fields.docx",
    )
    s1 = DocSection(doc)
    s1.next_section()
    s2 = s1.deeper()
    s2.next_section()
    s3 = s2.deeper()
    s3.next_section()
    assert s3.first_paragraph_text() == "<software-component.description>"
    s3.replace_first_paragraph(None)
    replacement = "Replacement text"
    s3.add_run(replacement)
    assert s3.first_paragraph_text() == replacement
    s2 = s1.deeper()
    s2.next_section()
    s3 = s2.deeper()
    s3.next_section()
    assert s3.first_paragraph_text() == replacement
    assert s2.next_section()
    assert s2.first_paragraph_text() == "Heading Two"


def test_docsection_replace_field():
    """Test field replacement."""
    doc = docx.Document(
        "tests/resources/msword/DocSection_fields.docx",
    )
    s1 = DocSection(doc)
    s1.next_section()
    s2 = s1.deeper()
    s2.next_section()
    component_id = "software-component.id"
    component_name = "software-component.name"
    assert s2.fields() == {
        component_id,
        component_name,
    }
    s2.replace_field(component_name, "The Name")
    assert s2.fields() == {
        component_id,
    }
    assert s2.first_paragraph_text() == (
        f"Paragraph with fields <{component_id}> and The Name."
    )
